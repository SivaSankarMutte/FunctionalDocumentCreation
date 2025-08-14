import sys, traceback
sys.excepthook = lambda exctype, value, tb: traceback.print_exception(exctype, value, tb)


import os
import io
import re
import zipfile
import hashlib
from pathlib import Path
from typing import List, Tuple

from dotenv import load_dotenv
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_core.documents import Document
from langchain_core.prompts import ChatPromptTemplate
from langchain_groq import ChatGroq
from langchain_openai import ChatOpenAI
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_chroma import Chroma

from docx import Document as DocxDocument

# ---------- Config (for Local) ----------
# load_dotenv()
# EMBED_MODEL = os.getenv("EMBED_MODEL", "sentence-transformers/all-MiniLM-L6-v2")
# LLM_MODEL = os.getenv("LLM_MODEL")  # optional; choose based on provider
# PERSIST_DIR = os.getenv("PERSIST_DIR", "./code_index")
# CHUNK_SIZE = int(os.getenv("CHUNK_SIZE", "1200"))
# CHUNK_OVERLAP = int(os.getenv("CHUNK_OVERLAP", "200"))

# ---------- Config (for Streamlit) ----------
EMBED_MODEL = st.secrets.get("EMBED_MODEL", "sentence-transformers/paraphrase-MiniLM-L3-v2")
LLM_MODEL = st.secrets.get("LLM_MODEL", None)  # optional; choose based on provider
PERSIST_DIR = st.secrets.get("PERSIST_DIR", "./code_index")
CHUNK_SIZE = int(st.secrets.get("CHUNK_SIZE", 1200))
CHUNK_OVERLAP = int(st.secrets.get("CHUNK_OVERLAP", 200))

MAX_FILE_SIZE_MB = 2.5  # skip very large single files
MAX_DOCS = 5000  # safety cap

IGNORES = [
    r"/\.git/", r"/\.hg/", r"/\.svn/",
    r"/node_modules/", r"/dist/", r"/build/", r"/coverage/", r"/.next/", r"/.nuxt/",
    r"/bin/", r"/obj/", r"/target/", r"/out/",
    r"/venv/", r"/\.venv/", r"/\.pytest_cache/",
    r"/__pycache__/", r"/\.DS_Store$",
]
BINARY_EXTS = [
    ".png",".jpg",".jpeg",".gif",".bmp",".ico",".pdf",".zip",".jar",".exe",".dll",".so",".dylib",".ttf",".woff",".woff2",".eot",".mp4",".mp3"
]

# ---------- Helpers ----------
def _should_ignore(path: str) -> bool:
    for pat in IGNORES:
        if re.search(pat, path.replace("\\", "/")):
            return True
    return False

def _is_text_file(path: str) -> bool:
    return os.path.splitext(path)[1].lower() not in BINARY_EXTS

def _hash_text(t: str) -> str:
    return hashlib.sha1(t.encode("utf-8")).hexdigest()[:10]

# ---------- Ingest ZIP ----------
def extract_zip_to_temp(zip_bytes: bytes, dest: str) -> str:
    os.makedirs(dest, exist_ok=True)
    with zipfile.ZipFile(io.BytesIO(zip_bytes)) as zf:
        zf.extractall(dest)
    return dest

def collect_code_files(root: str) -> List[str]:
    files = []
    for dirpath, _, filenames in os.walk(root):
        if _should_ignore(dirpath):
            continue
        for fn in filenames:
            fp = os.path.join(dirpath, fn)
            if _should_ignore(fp): 
                continue
            if not _is_text_file(fp):
                continue
            try:
                if os.path.getsize(fp) > MAX_FILE_SIZE_MB * 1024 * 1024:
                    continue
                files.append(fp)
            except FileNotFoundError:
                continue
    return files

def read_file(path: str) -> str:
    try:
        with open(path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()
    except Exception as e:
        return ""

# ---------- Build Index ----------
def build_index_from_folder(folder: str) -> Chroma:
    files = collect_code_files(folder)
    if not files:
        raise ValueError("No code files found to index.")

    docs = []
    for fp in files:
        content = read_file(fp)
        if not content.strip():
            continue
        meta = {"source": fp}
        docs.append(Document(page_content=content, metadata=meta))
        if len(docs) >= MAX_DOCS:
            break

    splitter = RecursiveCharacterTextSplitter(chunk_size=CHUNK_SIZE, chunk_overlap=CHUNK_OVERLAP)
    chunks = splitter.split_documents(docs)

    embeddings = HuggingFaceEmbeddings(model_name=EMBED_MODEL, model_kwargs={"device":"cpu"})
    vs = Chroma.from_documents(chunks, embedding=embeddings, persist_directory=PERSIST_DIR)
    return vs

def load_or_build_index(zip_bytes: bytes, workdir: str) -> Chroma:
    project_root = extract_zip_to_temp(zip_bytes, os.path.join(workdir, "repo"))
    # Always rebuild for this demo to capture the uploaded repo
    if os.path.exists(PERSIST_DIR):
        # clear old index
        for f in os.listdir(PERSIST_DIR):
            fp = os.path.join(PERSIST_DIR, f)
            if os.path.isdir(fp):
                import shutil; shutil.rmtree(fp, ignore_errors=True)
            else:
                os.remove(fp)
    return build_index_from_folder(project_root)

# ---------- LLM (for Local) ----------
# def get_llm():
#     if os.getenv("GROQ_API_KEY"):
#         return ChatGroq(model_name=LLM_MODEL or "llama3-70b-8192", temperature=0.2)
#     if os.getenv("OPENAI_API_KEY"):
#         return ChatOpenAI(model=LLM_MODEL or "gpt-4o-mini", temperature=0.2)
#     raise RuntimeError("Please set GROQ_API_KEY or OPENAI_API_KEY.")

# ---------- LLM (for Streamlit) ----------
def get_llm():
    if st.secrets.get("GROQ_API_KEY"):
        return ChatGroq(model_name=LLM_MODEL or "llama3-70b-8192", temperature=0.2)
    if st.secrets.get("OPENAI_API_KEY"):
        return ChatOpenAI(model=LLM_MODEL or "gpt-4o-mini", temperature=0.2)
    raise RuntimeError("Please set GROQ_API_KEY or OPENAI_API_KEY.")

# ---------- Document Generation ----------
SECTION_QUERIES = [
    ("Project Overview & Purpose",
     "Summarize the business purpose and high-level functionality of this application based on the codebase. Include main domains and user roles."),
    ("Feature Modules & Responsibilities",
     "List major modules/components and describe their responsibilities. Mention key files and directories for each."),
    ("UI Screens & Routes",
     "If the project has a UI, identify main screens or routes and describe the user journey. Reference React routes, Razor pages, or controllers."),
    ("API Endpoints",
     "Identify API endpoints (HTTP method, path) and map them to handlers/controllers. Describe inputs, outputs, and side effects if visible."),
    ("Data & Control Flow",
     "Describe data flow across layers (UI → API → Services → DB). Include notable design patterns (e.g., repository, CQRS, Redux)."),
    ("Configuration & Environment",
     "List environment variables, configuration files, and secrets usage. Explain how builds and environments are configured."),
    ("Dependencies & Integrations",
     "List key third-party libraries and external services, and explain what they’re used for."),
    ("Build, Run & Deploy",
     "Explain how to set up, build, test, and deploy the application based on scripts and config files."),
    ("Risks, TODOs & Tech Debt",
     "Identify potential risks, TODO comments, deprecated code, and areas needing documentation or tests."),
]

def generate_section(llm, retriever, title: str, question: str) -> str:
    prompt = ChatPromptTemplate.from_template(
        """You are a senior software analyst.
Use the retrieved code snippets as ground truth to answer the question.
Be concise, structured, and avoid guessing. Cite file paths where relevant.

# Question
{question}

# Retrieved Context
{context}

# Answer
"""
    )
    # Pull top-k chunks for the section
    docs = retriever.get_relevant_documents(question)
    context = "\n\n---\n\n".join([f"[{d.metadata.get('source','?')}]\\n{d.page_content[:2000]}" for d in docs])
    chain = prompt | llm
    ans = chain.invoke({"question": question, "context": context}).content
    return f"## {title}\n\n{ans.strip()}\n"

def generate_functional_doc(zip_bytes: bytes, workdir: str = "./work") -> Tuple[str, str]:
    os.makedirs(workdir, exist_ok=True)
    vs = load_or_build_index(zip_bytes, workdir)
    retriever = vs.as_retriever(search_kwargs={"k": 12})
    llm = get_llm()

    # Build markdown
    md_parts = ["# Functional Documentation\n"]
    for title, q in SECTION_QUERIES:
        md_parts.append(generate_section(llm, retriever, title, q))
    md = "\n".join(md_parts)

    out_dir = os.path.join(workdir, "output")
    os.makedirs(out_dir, exist_ok=True)
    md_path = os.path.join(out_dir, "functional_doc.md")
    with open(md_path, "w", encoding="utf-8") as f:
        f.write(md)

    # Also write DOCX
    doc = DocxDocument()
    doc.add_heading("Functional Documentation", level=0)
    for title, _ in SECTION_QUERIES:
        # Extract section content
        marker = f"## {title}"
        start = md.find(marker)
        end = md.find("## ", start + 3) if start != -1 else -1
        section_text = md[start:end].replace(marker, title).strip() if start != -1 else ""
        doc.add_heading(title, level=1)
        for para in section_text.split("\\n\\n"):
            if para.strip():
                doc.add_paragraph(para)
    docx_path = os.path.join(out_dir, "functional_doc.docx")
    doc.save(docx_path)

    return md_path, docx_path
